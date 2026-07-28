import docx
from pathlib import Path
from app.parser.text_cleaner import clean_text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a .docx file and return cleaned string.
    """
    docx_file = Path(docx_path)

    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        doc = docx.Document(str(docx_file))
        full_text = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        raw_text = "\n".join(full_text)
        return clean_text(raw_text)

    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")
